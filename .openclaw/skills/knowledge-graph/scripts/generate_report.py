#!/usr/bin/env python3
"""
Report Generation - Generate Markdown or Word reports from video analysis results

Usage:
    python3 generate_report.py --video-info 'JSON' --analysis 'JSON' [--format markdown]
    python3 generate_report.py --video-info-file video.json --analysis-file analysis.json [--format word]
"""
import argparse
import json
import logging
import sys
from datetime import datetime
from pathlib import Path

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s", stream=sys.stderr)
logger = logging.getLogger(__name__)


def _load_json_arg(value: str, file_value: str = None):
    if file_value:
        return json.loads(Path(file_value).read_text(encoding="utf-8"))
    return json.loads(value)


def sanitize_filename(name: str) -> str:
    for char in '<>:"/\\|?*':
        name = name.replace(char, "_")
    return name[:100]


def format_duration(seconds: int) -> str:
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60
    if hours > 0:
        return f"{hours}h {minutes}m {secs}s"
    elif minutes > 0:
        return f"{minutes}m {secs}s"
    return f"{secs}s"


def format_timestamp(seconds: float) -> str:
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def build_markdown(video_info: dict, analysis: dict) -> str:
    lines = []

    # Title
    lines.append(f"# {video_info.get('title', 'Untitled Video')}\n")

    # Video info
    lines.append("## Video Info\n")
    lines.append(f"- **Source**: {video_info.get('platform', 'unknown')}")
    lines.append(f"- **Link**: {video_info.get('url', '')}")
    lines.append(f"- **Duration**: {format_duration(video_info.get('duration', 0))}")
    lines.append(f"- **Generated**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

    # Summary
    summary = analysis.get("summary", "")
    if summary:
        lines.append("## Summary\n")
        lines.append(f"{summary}\n")

    # Key points
    key_points = analysis.get("key_points", [])
    if key_points:
        lines.append("## Key Points\n")
        for point in key_points:
            lines.append(f"- {point}")
        lines.append("")

    # Transcript
    transcript = analysis.get("transcript", "")
    if transcript:
        lines.append("## Transcript\n")
        lines.append(f"```\n{transcript[:3000]}...\n```\n")

    # Entities
    entities = analysis.get("entities", [])
    if entities:
        lines.append("## Entity Recognition\n")
        lines.append(f"Identified {len(entities)} entities:\n")
        entity_groups = {}
        for entity in entities:
            label = entity.get("label", "UNKNOWN")
            if label not in entity_groups:
                entity_groups[label] = []
            entity_groups[label].append(entity.get("text", ""))
        for label, items in entity_groups.items():
            lines.append(f"- **{label}**: {', '.join(items)}")
        lines.append("")

    # Relations
    relations = analysis.get("relations", [])
    if relations:
        lines.append("## Entity Relations\n")
        lines.append(f"Identified {len(relations)} relations:\n")
        for rel in relations[:20]:
            lines.append(f"- {rel.get('source', '')} -> {rel.get('target', '')} ({rel.get('relation', '')})")
        lines.append("")

    # Timeline
    if "timeline" in analysis:
        lines.append("## Timeline\n")
        for item in analysis["timeline"]:
            ts = item.get("timestamp", 0)
            text = item.get("text", "")
            lines.append(f"- **{format_timestamp(ts)}**: {text[:100]}...")
        lines.append("")

    return "\n".join(lines)


def generate_word(video_info: dict, analysis: dict, output_path: Path):
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        logger.error("python-docx not installed, cannot generate Word document")
        return None

    doc = Document()
    doc.add_heading(video_info.get("title", "Untitled Video"), 0)

    doc.add_heading("Video Info", level=1)
    p = doc.add_paragraph()
    p.add_run("Source: ").bold = True
    p.add_run(f"{video_info.get('platform', 'unknown')}\n")
    p.add_run("Link: ").bold = True
    p.add_run(f"{video_info.get('url', '')}\n")
    p.add_run("Duration: ").bold = True
    p.add_run(f"{format_duration(video_info.get('duration', 0))}\n")
    p.add_run("Generated: ").bold = True
    p.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

    summary = analysis.get("summary", "")
    if summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

    transcript = analysis.get("transcript", "")
    if transcript:
        doc.add_heading("Transcript", level=1)
        doc.add_paragraph(transcript[:2000])

    entities = analysis.get("entities", [])
    if entities:
        doc.add_heading("Entity Recognition", level=1)
        doc.add_paragraph(f"Identified {len(entities)} entities")
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Type"
        hdr[1].text = "Entity"
        for entity in entities:
            row = table.add_row().cells
            row[0].text = entity.get("label", "")
            row[1].text = entity.get("text", "")

    doc.save(output_path)
    return str(output_path)


def generate_report(video_info: dict, analysis: dict, fmt: str = "markdown", output_dir: str = "data/docs") -> dict:
    out_path = Path(output_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    safe_title = sanitize_filename(video_info.get("title", "untitled"))
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"{safe_title}_{timestamp}"

    doc_path = None

    if fmt in ("markdown", "both"):
        md_path = out_path / f"{base_name}.md"
        content = build_markdown(video_info, analysis)
        md_path.write_text(content, encoding="utf-8")
        doc_path = str(md_path)
        logger.info("Generated Markdown: %s", md_path)

    if fmt in ("word", "both"):
        docx_path = out_path / f"{base_name}.docx"
        result = generate_word(video_info, analysis, docx_path)
        if result:
            if not doc_path:
                doc_path = result
            logger.info("Generated Word: %s", result)

    return {"document_path": doc_path, "format": fmt}


def main():
    parser = argparse.ArgumentParser(description="Generate Markdown or Word reports from video analysis results")
    parser.add_argument("--video-info", help="Video info JSON string")
    parser.add_argument("--video-info-file", help="Video info JSON file path")
    parser.add_argument("--analysis", help="Analysis result JSON string")
    parser.add_argument("--analysis-file", help="Analysis result JSON file path")
    parser.add_argument("--format", choices=["markdown", "word", "both"], default="markdown", help="Output format (default: markdown)")
    parser.add_argument("--output-dir", default="data/docs", help="Output directory (default: data/docs)")
    args = parser.parse_args()

    video_info = _load_json_arg(args.video_info or "{}", args.video_info_file)
    analysis = _load_json_arg(args.analysis or "{}", args.analysis_file)

    result = generate_report(video_info, analysis, fmt=args.format, output_dir=args.output_dir)
    json.dump(result, sys.stdout, ensure_ascii=False, indent=2)
    print()


if __name__ == "__main__":
    main()
