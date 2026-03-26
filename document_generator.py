"""
文档生成模块 - Markdown/Word
"""
import logging
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from config import DOC_DIR, SUMMARY_MAX_LENGTH

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentGenerator:
    """文档生成器"""

    def __init__(self, output_format: str = "markdown"):
        self.output_format = output_format
        self.output_dir = DOC_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(self, video_info: Dict[str, Any], analysis: Dict[str, Any]) -> str:
        """
        生成完整的分析文档

        Args:
            video_info: 视频信息（从VideoExtractor获取）
            analysis: 分析结果（包含transcript、entities、relations等）

        Returns:
            生成文档的路径
        """
        # 生成文件名
        safe_title = self._sanitize_filename(video_info.get("title", "untitled"))
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = f"{safe_title}_{timestamp}"

        # 生成Markdown
        md_path = self.output_dir / f"{base_name}.md"
        self._generate_markdown(md_path, video_info, analysis)

        # 可选：生成Word文档
        docx_path = None
        if self.output_format in ["word", "both"] and DOCX_AVAILABLE:
            docx_path = self.output_dir / f"{base_name}.docx"
            self._generate_word(docx_path, video_info, analysis)

        logger.info(f"Generated document: {md_path}")
        if docx_path:
            logger.info(f"Generated Word document: {docx_path}")

        return str(md_path)

    def _generate_markdown(self, path: Path, video_info: Dict, analysis: Dict):
        """生成Markdown文档"""
        content = self._build_markdown_content(video_info, analysis)

        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

    def _build_markdown_content(self, video_info: Dict, analysis: Dict) -> str:
        """构建Markdown内容"""
        lines = []

        # 标题
        lines.append(f"# {video_info.get('title', '未命名视频')}\n")

        # 元信息
        lines.append("## 视频信息\n")
        lines.append(f"- **来源**: {video_info.get('platform', 'unknown')}")
        lines.append(f"- **链接**: {video_info.get('url', '')}")
        lines.append(f"- **时长**: {self._format_duration(video_info.get('duration', 0))}")
        lines.append(f"- **生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

        # 摘要
        summary = analysis.get("summary", "")
        if summary:
            lines.append("## 摘要\n")
            lines.append(f"{summary[:SUMMARY_MAX_LENGTH]}...\n")

        # 转录文本
        transcript = analysis.get("transcript", "")
        if transcript:
            lines.append("## 转录文本\n")
            lines.append(f"```\n{transcript[:3000]}...\n```\n")

        # 实体
        entities = analysis.get("entities", [])
        if entities:
            lines.append("## 实体识别\n")
            lines.append(f"共识别 {len(entities)} 个实体：\n")

            # 按类型分组
            entity_groups = {}
            for entity in entities:
                label = entity.get("label", "UNKNOWN")
                if label not in entity_groups:
                    entity_groups[label] = []
                entity_groups[label].append(entity.get("text", ""))

            for label, items in entity_groups.items():
                lines.append(f"- **{label}**: {', '.join(items)}")
            lines.append("")

        # 关系
        relations = analysis.get("relations", [])
        if relations:
            lines.append("## 实体关系\n")
            lines.append(f"共识别 {len(relations)} 条关系：\n")
            for rel in relations[:20]:  # 限制显示数量
                source = rel.get("source", "")
                target = rel.get("target", "")
                rel_type = rel.get("relation", "")
                lines.append(f"- {source} → {target} ({rel_type})")
            lines.append("")

        # 时间线
        if "timeline" in analysis:
            lines.append("## 时间线\n")
            for item in analysis["timeline"]:
                timestamp = item.get("timestamp", 0)
                text = item.get("text", "")
                lines.append(f"- **{self._format_timestamp(timestamp)}**: {text[:100]}...")
            lines.append("")

        # 相关视频
        if "related_videos" in analysis:
            lines.append("## 相关内容\n")
            for related in analysis["related_videos"]:
                lines.append(f"- {related}\n")

        return "\n".join(lines)

    def _generate_word(self, path: Path, video_info: Dict, analysis: Dict):
        """生成Word文档"""
        doc = Document()

        # 标题
        doc.add_heading(video_info.get("title", "未命名视频"), 0)

        # 元信息
        doc.add_heading("视频信息", level=1)
        p = doc.add_paragraph()
        p.add_run("来源: ").bold = True
        p.add_run(f"{video_info.get('platform', 'unknown')}\n")
        p.add_run("链接: ").bold = True
        p.add_run(f"{video_info.get('url', '')}\n")
        p.add_run("时长: ").bold = True
        p.add_run(f"{self._format_duration(video_info.get('duration', 0))}\n")
        p.add_run("生成时间: ").bold = True
        p.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

        # 摘要
        summary = analysis.get("summary", "")
        if summary:
            doc.add_heading("摘要", level=1)
            doc.add_paragraph(summary[:SUMMARY_MAX_LENGTH])

        # 转录文本
        transcript = analysis.get("transcript", "")
        if transcript:
            doc.add_heading("转录文本", level=1)
            doc.add_paragraph(transcript[:2000])

        # 实体
        entities = analysis.get("entities", [])
        if entities:
            doc.add_heading("实体识别", level=1)
            doc.add_paragraph(f"共识别 {len(entities)} 个实体")

            # 表格
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "类型"
            hdr_cells[1].text = "实体名"

            for entity in entities:
                row_cells = table.add_row().cells
                row_cells[0].text = entity.get("label", "")
                row_cells[1].text = entity.get("text", "")

        # 保存
        doc.save(path)

    def generate_summary(self, text: str, max_length: int = SUMMARY_MAX_LENGTH) -> str:
        """生成摘要（简单截取）"""
        # TODO: 实际应该用LLM生成摘要
        sentences = text.split("。")
        summary = ""
        for sent in sentences:
            if len(summary) + len(sent) > max_length:
                break
            summary += sent + "。"

        return summary or text[:max_length]

    @staticmethod
    def _sanitize_filename(name: str) -> str:
        """清理文件名"""
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            name = name.replace(char, "_")
        return name[:100]  # 限制长度

    @staticmethod
    def _format_duration(seconds: int) -> str:
        """格式化时长"""
        hours = seconds // 3600
        minutes = (seconds % 3600) // 60
        secs = seconds % 60

        if hours > 0:
            return f"{hours}h {minutes}m {secs}s"
        elif minutes > 0:
            return f"{minutes}m {secs}s"
        else:
            return f"{secs}s"

    @staticmethod
    def _format_timestamp(seconds: float) -> str:
        """格式化时间戳"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"


# CLI测试
if __name__ == "__main__":
    import sys

    # 模拟数据
    video_info = {
        "title": "测试视频标题",
        "platform": "youtube",
        "url": "https://youtube.com/watch?v=test",
        "duration": 600
    }

    analysis = {
        "summary": "这是一个关于人工智能的测试视频。",
        "transcript": "张三是北京大学的教授。他专注于机器学习研究。中国计算机学会给了他奖项。",
        "entities": [
            {"text": "张三", "label": "PERSON"},
            {"text": "北京大学", "label": "ORG"},
            {"text": "机器学习", "label": "CONCEPT"}
        ],
        "relations": []
    }

    gen = DocumentGenerator(output_format="markdown")
    path = gen.generate(video_info, analysis)

    print(f"Generated: {path}")
