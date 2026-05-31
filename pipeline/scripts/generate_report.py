#!/usr/bin/env python3
"""
报告生成 — 从视频分析结果生成 Markdown 或 Word 报告

Usage:
    python3 generate_report.py --video-info 'JSON' --analysis 'JSON' [--format markdown]
    python3 generate_report.py --video-info-file video.json --analysis-file analysis.json [--format word]
"""
import argparse
import json
import logging
import sys
from datetime import datetime
from pathlib import Path

logging.basicConfig(level=logging.INFO, format="[%(levelname)s] %(message)s", stream=sys.stderr)
logger = logging.getLogger(__name__)


def _load_json_arg(value: str, file_value: str = None):
    if file_value:
        return json.loads(Path(file_value).read_text(encoding="utf-8"))
    return json.loads(value)


def sanitize_filename(name: str) -> str:
    for char in '<>:"/\\|?*':
        name = name.replace(char, "_")
    return name[:100]


def format_duration(seconds: int) -> str:
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60
    if hours > 0:
        return f"{hours}h {minutes}m {secs}s"
    elif minutes > 0:
        return f"{minutes}m {secs}s"
    return f"{secs}s"


def format_timestamp(seconds: float) -> str:
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def build_markdown(video_info: dict, analysis: dict) -> str:
    lines = []

    # 标题
    lines.append(f"# {video_info.get('title', 'Untitled Video')}\n")

    # 视频信息
    lines.append("## 视频信息\n")
    lines.append(f"- **来源**: {video_info.get('platform', 'unknown')}")
    lines.append(f"- **链接**: {video_info.get('url', '')}")
    lines.append(f"- **时长**: {format_duration(video_info.get('duration', 0))}")
    lines.append(f"- **生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

    # 摘要 — 使用 full_summary 字段（对齐 schema.py）
    summary = analysis.get("full_summary", "") or analysis.get("summary", "")
    if summary:
        lines.append("## 摘要\n")
        lines.append(f"{summary}\n")

    # 要点
    key_points = analysis.get("key_points", [])
    if key_points:
        lines.append("## 关键要点\n")
        for point in key_points:
            lines.append(f"- {point}")
        lines.append("")

    # 转写文本
    transcript = analysis.get("transcript", "")
    if transcript:
        lines.append("## 转写文本\n")
        lines.append(f"```\n{transcript[:3000]}...\n```\n")

    # 实体 — 使用 name 字段（对齐 schema.py）
    entities = analysis.get("entities", [])
    if entities:
        lines.append("## 实体识别\n")
        lines.append(f"识别出 {len(entities)} 个实体:\n")
        entity_groups = {}
        for entity in entities:
            etype = entity.get("type", entity.get("label", "UNKNOWN"))
            ename = entity.get("name", entity.get("text", ""))
            if etype not in entity_groups:
                entity_groups[etype] = []
            entity_groups[etype].append(ename)
        for etype, items in entity_groups.items():
            lines.append(f"- **{etype}**: {', '.join(items)}")
        lines.append("")

    # 关系
    relations = analysis.get("relations", [])
    if relations:
        lines.append("## 实体关系\n")
        lines.append(f"识别出 {len(relations)} 个关系:\n")
        for rel in relations[:20]:
            lines.append(f"- {rel.get('source', '')} → {rel.get('target', '')} ({rel.get('relation', '')})")
        lines.append("")

    # 时间线
    if "timeline" in analysis:
        lines.append("## 时间线\n")
        for item in analysis["timeline"]:
            ts = item.get("timestamp", 0)
            text = item.get("text", "")
            lines.append(f"- **{format_timestamp(ts)}**: {text[:100]}...")
        lines.append("")

    return "\n".join(lines)


def generate_word(video_info: dict, analysis: dict, output_path: Path):
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        logger.error("python-docx 未安装，无法生成 Word 文档")
        return None

    doc = Document()
    doc.add_heading(video_info.get("title", "Untitled Video"), 0)

    doc.add_heading("视频信息", level=1)
    p = doc.add_paragraph()
    p.add_run("来源: ").bold = True
    p.add_run(f"{video_info.get('platform', 'unknown')}\n")
    p.add_run("链接: ").bold = True
    p.add_run(f"{video_info.get('url', '')}\n")
    p.add_run("时长: ").bold = True
    p.add_run(f"{format_duration(video_info.get('duration', 0))}\n")
    p.add_run("生成时间: ").bold = True
    p.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

    summary = analysis.get("full_summary", "") or analysis.get("summary", "")
    if summary:
        doc.add_heading("摘要", level=1)
        doc.add_paragraph(summary)

    transcript = analysis.get("transcript", "")
    if transcript:
        doc.add_heading("转写文本", level=1)
        doc.add_paragraph(transcript[:2000])

    entities = analysis.get("entities", [])
    if entities:
        doc.add_heading("实体识别", level=1)
        doc.add_paragraph(f"识别出 {len(entities)} 个实体")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "类型"
        hdr[1].text = "实体名称"
        hdr[2].text = "描述"
        for entity in entities:
            row = table.add_row().cells
            row[0].text = entity.get("type", entity.get("label", ""))
            row[1].text = entity.get("name", entity.get("text", ""))
            row[2].text = entity.get("description", entity.get("context", ""))

    doc.save(output_path)
    return str(output_path)


def generate_report(video_info: dict, analysis: dict, fmt: str = "markdown", output_dir: str = "data/docs") -> dict:
    out_path = Path(output_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    safe_title = sanitize_filename(video_info.get("title", "untitled"))
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"{safe_title}_{timestamp}"

    doc_path = None

    if fmt in ("markdown", "both"):
        md_path = out_path / f"{base_name}.md"
        content = build_markdown(video_info, analysis)
        md_path.write_text(content, encoding="utf-8")
        doc_path = str(md_path)
        logger.info("已生成 Markdown: %s", md_path)

    if fmt in ("word", "both"):
        docx_path = out_path / f"{base_name}.docx"
        result = generate_word(video_info, analysis, docx_path)
        if result:
            if not doc_path:
                doc_path = result
            logger.info("已生成 Word: %s", result)

    return {"document_path": doc_path, "format": fmt}


def main():
    parser = argparse.ArgumentParser(description="从视频分析结果生成 Markdown 或 Word 报告")
    parser.add_argument("--video-info", help="视频信息 JSON 字符串")
    parser.add_argument("--video-info-file", help="视频信息 JSON 文件路径")
    parser.add_argument("--analysis", help="分析结果 JSON 字符串")
    parser.add_argument("--analysis-file", help="分析结果 JSON 文件路径")
    parser.add_argument("--format", choices=["markdown", "word", "both"], default="markdown", help="输出格式 (默认: markdown)")
    parser.add_argument("--output-dir", default="data/docs", help="输出目录 (默认: data/docs)")
    args = parser.parse_args()

    video_info = _load_json_arg(args.video_info or "{}", args.video_info_file)
    analysis = _load_json_arg(args.analysis or "{}", args.analysis_file)

    result = generate_report(video_info, analysis, fmt=args.format, output_dir=args.output_dir)
    json.dump(result, sys.stdout, ensure_ascii=False, indent=2)
    print()


if __name__ == "__main__":
    main()
